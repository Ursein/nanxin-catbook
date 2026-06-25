"""
Convert Markdown to .docx with proper formatting.
Usage: py scripts/md2docx.py <input.md> <output.docx>
"""
import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)


def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    list_items = []  # buffer for consecutive list items
    list_ordered = False
    list_level = 0

    def flush_list():
        nonlocal list_items, list_ordered, list_level
        if not list_items:
            return
        for idx, (text, level) in enumerate(list_items):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(text)
            _apply_inline_formatting(run, text)
        list_items = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        # Parse table
        rows_data = []
        for line in table_lines:
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                rows_data.append(cells)

        if len(rows_data) < 2:
            table_lines = []
            return

        # Separate header and data
        header = rows_data[0]
        # Skip separator row (|---|)
        data_rows = [r for r in rows_data[1:] if not all(re.match(r'^[-:]+$', c) for c in r)]

        num_cols = len(header)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_borders(table)

        # Header row
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D9E2F3')

        # Data rows
        for ri, row_data in enumerate(data_rows):
            for j, cell_text in enumerate(row_data):
                if j >= num_cols:
                    break
                cell = table.rows[ri + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                # Alternate row shading
                if ri % 2 == 0:
                    set_cell_shading(cell, 'F2F2F2')

        # Set column widths - proportional to page width
        page_width = 6.3  # inches (A4/letter minus margins)
        header_words = [h.strip() for h in header]

        # Smart column distribution
        if num_cols == 2 and header_words[0] in ('项目', '编号', '功能编号'):
            # Key-value tables: key column narrower, value column wide
            col_widths = [Inches(1.2), Inches(page_width - 1.2)]
        elif num_cols == 3:
            # Three-column: distribute evenly but wider
            col_widths = [Inches(page_width / 3)] * num_cols
        elif num_cols == 4:
            col_widths = [Inches(page_width / 4)] * num_cols
        elif num_cols >= 5:
            # Many columns: each gets reasonable width, horizontal scroll in Word
            col_widths = [Inches(min(page_width / num_cols, 1.8)) for _ in range(num_cols)]
            # Make first column slightly narrower for IDs
            col_widths[0] = Inches(0.8)
        else:
            col_widths = [Inches(page_width / num_cols) for _ in range(num_cols)]

        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(col_widths):
                    cell.width = col_widths[idx]

        doc.add_paragraph()  # spacing after table
        table_lines = []

    def _apply_inline_formatting(run, text):
        """Apply bold/italic formatting detected in text."""
        # If no formatting markers, just set text
        if '**' not in text and '*' not in text.replace('**', ''):
            run.text = text
            return

        # Simple approach: replace **text** with bold runs
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            else:
                run.text = part

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('\n'.join(code_lines) if code_lines else '(code block)')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
                code_lines = []
                in_code_block = False
                doc.add_paragraph()
            else:
                flush_list()
                flush_table()
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            flush_list()
            flush_table()
            # Skip — no need for separator in docx
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            flush_list()
            flush_table()
            i += 1
            continue

        # Table row
        if line.strip().startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        else:
            flush_table()

        # Heading detection: # heading
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line.strip())
        if heading_match:
            flush_list()
            flush_table()
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level <= 4:
                p = doc.add_heading(text, level=level)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line.strip())
        if ol_match:
            indent = len(ol_match.group(1)) // 4
            text = ol_match.group(3)
            list_items.append((text, indent))
            list_ordered = True
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^(\s*)[-*]\s+(.+)$', line.strip())
        if ul_match:
            indent = len(ul_match.group(1)) // 4
            text = ul_match.group(2)
            list_items.append((text, indent))
            list_ordered = False
            i += 1
            continue

        flush_list()
        flush_table()

        # Regular paragraph - handle inline formatting
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            # Check for bold markers in the text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)

        i += 1

    # Flush remaining
    flush_list()
    flush_table()

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: py scripts/md2docx.py <input.md> <output.docx>')
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    result = parse_markdown_to_docx(md_path, docx_path)
    print(f'[OK] Generated: {result}')