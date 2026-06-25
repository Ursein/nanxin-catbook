"""
南信猫友记 - 按软件工程逻辑重组文档
结构：研究背景 → 可行性 → 需求分析 → 概要设计 → 详细实现 → 软件测试
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'C:\Users\16570\Desktop\nanxin-catbook\docs'

# ============================================================
# 1. 从 MD 文件中提取指定章节
# ============================================================
def extract_section(lines, start_pattern, end_patterns=None, include_heading=True):
    """提取从 start_pattern 匹配行开始，到 end_patterns 中任一匹配行之前的内容"""
    result = []
    capturing = False
    for i, line in enumerate(lines):
        if not capturing:
            if re.match(start_pattern, line.strip()):
                capturing = True
                if include_heading:
                    result.append(line)
                continue
        else:
            if end_patterns:
                for ep in end_patterns:
                    if re.match(ep, line.strip()):
                        return result
            else:
                # 遇到同级或更高级标题就停止
                if re.match(r'^##\s+', line.strip()):
                    return result
            result.append(line)
    return result

def extract_between(lines, start_pattern, end_pattern):
    """提取两个标记之间的内容（不含边界）"""
    result = []
    capturing = False
    for line in lines:
        if not capturing:
            if re.match(start_pattern, line.strip()):
                capturing = True
                continue
        else:
            if re.match(end_pattern, line.strip()):
                return result
            result.append(line)
    return result

def get_section_text(lines, start_pattern, end_pattern=None):
    """从 lines 中提取从 start 到 end 的内容，去掉首尾空行"""
    result = extract_between(lines, start_pattern, end_pattern) if end_pattern else extract_section(lines, start_pattern)
    while result and not result[0].strip():
        result.pop(0)
    while result and not result[-1].strip():
        result.pop()
    return result

def shift_headings(lines, levels=1):
    """将所有标题降一级（多加 levels 个 #）"""
    result = []
    for line in lines:
        m = re.match(r'^(#{1,6})\s', line)
        if m:
            hashes = m.group(1)
            new_hashes = hashes + '#' * levels
            if len(new_hashes) <= 6:
                line = new_hashes + line[len(hashes):]
        result.append(line)
    return result

def read_md_lines(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read().split('\n')

# ============================================================
# 2. 自然段改写
# ============================================================
def bullets_to_paragraph(lines, category_prefix=""):
    sentences = []
    for line in lines:
        text = line.strip().lstrip('-').lstrip('•').strip()
        if text:
            text = text.rstrip('。').rstrip('.').strip()
            sentences.append(text)
    if not sentences:
        return ""
    if len(sentences) == 1:
        return f"{category_prefix}{sentences[0]}。"
    if len(sentences) == 2:
        return f"{category_prefix}{sentences[0]}，{sentences[1]}。"
    return f"{category_prefix}{sentences[0]}，{'，'.join(sentences[1:-1])}，{sentences[-1]}。"

def convert_bullet_sections(text):
    lines = text.split('\n')
    result = []
    i = 0
    in_bullet_group = False
    bullet_lines = []
    bullet_context = ""
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 先检测需要保留上下文的特殊标题（安全性/可用性/可维护性等）
        if re.match(r'^#{3,5}\s*(安全性|可用性|可维护性|性能|约束|条件|限制|结论)', stripped):
            if in_bullet_group and bullet_lines:
                para = bullets_to_paragraph(bullet_lines, bullet_context)
                result.append('')
                result.append(para)
                result.append('')
                bullet_lines = []
            result.append(line)
            title_text = re.sub(r'^#+\s*', '', stripped)
            if title_text.endswith('需求') or title_text.endswith('性'):
                bullet_context = f"在{title_text}方面，"
            elif title_text:
                bullet_context = f"{title_text}方面，"
            in_bullet_group = False
            i += 1
            continue
        
        # 通用标题检测
        if re.match(r'^#{1,4}\s+\S', stripped):
            if in_bullet_group and bullet_lines:
                para = bullets_to_paragraph(bullet_lines, bullet_context)
                result.append('')
                result.append(para)
                result.append('')
                bullet_lines = []
                in_bullet_group = False
            result.append(line)
            bullet_context = ""
            i += 1
            continue
        
        # 旧的特殊标题检测已移到上面，此段删除
        
        if re.match(r'^[\s]*[-•]\s+', stripped):
            if not in_bullet_group:
                in_bullet_group = True
            bullet_lines.append(stripped)
            i += 1
            continue
        
        if in_bullet_group and bullet_lines:
            para = bullets_to_paragraph(bullet_lines, bullet_context)
            result.append('')
            result.append(para)
            result.append('')
            bullet_lines = []
            in_bullet_group = False
            bullet_context = ""
        
        result.append(line)
        i += 1
    
    if in_bullet_group and bullet_lines:
        para = bullets_to_paragraph(bullet_lines, bullet_context)
        result.append('')
        result.append(para)
        result.append('')
    
    return '\n'.join(result)

# ============================================================
# 3. 重组文档
# ============================================================
def build_restructured_md():
    req = read_md_lines(os.path.join(BASE, '需求分析', '需求规格说明书.md'))
    outline = read_md_lines(os.path.join(BASE, '概要设计', '概要设计说明书.md'))
    detail = read_md_lines(os.path.join(BASE, '详细设计', '详细设计说明书.md'))
    test = read_md_lines(os.path.join(BASE, '测试', '测试报告.md'))
    
    merged = []
    merged.append('# 南信猫友记——软件设计说明书\n')
    merged.append('> 南京信息工程大学 计算机与软件学院\n')
    merged.append('\n---\n')
    
    # ===== 第一章：研究背景、研究意义、研究现状 =====
    merged.append('## 第一章 研究背景与意义\n')
    merged.append('### 1.1 编写目的\n')
    merged += get_section_text(req, r'^### 1\.1 编写目的', r'^### 1\.2 研究背景')
    merged.append('')
    merged.append('### 1.2 研究背景\n')
    merged += get_section_text(req, r'^### 1\.2 研究背景', r'^### 1\.3 研究意义')
    merged.append('')
    merged.append('### 1.3 研究意义\n')
    merged += get_section_text(req, r'^### 1\.3 研究意义', r'^### 1\.4 研究现状')
    merged.append('')
    merged.append('### 1.4 研究现状\n')
    merged += get_section_text(req, r'^### 1\.4 研究现状', r'^### 1\.5 系统目标')
    merged.append('')
    merged.append('### 1.5 系统目标\n')
    merged += get_section_text(req, r'^### 1\.5 系统目标', r'^---')
    merged.append('')
    merged.append('---\n')
    
    # ===== 第二章：可行性分析 =====
    merged.append('## 第二章 可行性分析\n')
    merged.append('### 2.1 技术可行性\n')
    merged += get_section_text(req, r'^### 2\.1 技术可行性', r'^### 2\.2 经济可行性')
    merged.append('')
    merged.append('### 2.2 经济可行性\n')
    merged += get_section_text(req, r'^### 2\.2 经济可行性', r'^### 2\.3 操作可行性')
    merged.append('')
    merged.append('### 2.3 操作可行性\n')
    merged += get_section_text(req, r'^### 2\.3 操作可行性', r'^### 2\.4 法律')
    merged.append('')
    merged.append('### 2.4 法律可行性\n')
    merged += get_section_text(req, r'^### 2\.4 法律', r'^---')
    merged.append('')
    merged.append('---\n')
    
    # ===== 第三章：需求分析 =====
    merged.append('## 第三章 需求分析\n')
    merged.append('### 3.1 运行环境与限制\n')
    merged += get_section_text(req, r'^### 3\.1 运行环境', r'^### 3\.2 用例图')
    merged.append('')
    merged.append('### 3.2 用例图\n')
    merged += get_section_text(req, r'^### 3\.2 用例图', r'^### 3\.3 功能分析')
    merged.append('')
    merged.append('### 3.3 功能分析\n')
    merged += get_section_text(req, r'^### 3\.3 功能分析', r'^### 3\.4 性能分析')
    merged.append('')
    merged.append('### 3.4 性能分析\n')
    merged += get_section_text(req, r'^### 3\.4 性能分析', r'^### 3\.5 其他')
    merged.append('')
    merged.append('### 3.5 非功能需求\n')
    merged += get_section_text(req, r'^### 3\.5 其他非功能需求', r'^---')
    merged.append('')
    merged.append('---\n')
    
    # ===== 第四章：概要设计 =====
    merged.append('## 第四章 概要设计\n')
    merged.append('### 4.1 编写目的\n')
    merged += get_section_text(outline, r'^### 1\.1 编写目的', r'^### 1\.2 适用范围')
    merged.append('')
    merged.append('### 4.2 功能模块图\n')
    merged += shift_headings(get_section_text(outline, r'^## 2\. 功能模块图', r'^## 3\. 数据库'))
    merged.append('')
    merged.append('### 4.3 数据库 E-R 图\n')
    merged += shift_headings(get_section_text(outline, r'^## 3\. 数据库 E-R 图', r'^## 4\. 时序图'))
    merged.append('')
    merged.append('### 4.4 时序图\n')
    merged += shift_headings(get_section_text(outline, r'^## 4\. 时序图', r'^## 5\. 总结'))
    merged.append('')
    merged.append('---\n')
    
    # ===== 第五章：详细实现 =====
    merged.append('## 第五章 详细实现\n')
    merged.append('### 5.1 总体设计\n')
    merged += shift_headings(get_section_text(detail, r'^## 2\. 总体设计', r'^## 3\. 程序描述'))
    merged.append('')
    merged.append('### 5.2 程序流程与接口\n')
    merged += shift_headings(get_section_text(detail, r'^## 3\. 程序描述', r'^## 4\. 用户界面设计'))
    merged.append('')
    merged.append('### 5.3 用户界面设计\n')
    merged += shift_headings(get_section_text(detail, r'^## 4\. 用户界面设计', r'^## 5\. 总结'))
    merged.append('')
    merged.append('---\n')
    
    # ===== 第六章：软件测试 =====
    merged.append('## 第六章 软件测试\n')
    merged.append('### 6.1 测试目标与策略\n')
    merged += shift_headings(get_section_text(test, r'^## 2\. 测试计划', r'^## 3\. 单元测试'))
    merged.append('')
    merged.append('### 6.2 单元测试\n')
    merged += shift_headings(get_section_text(test, r'^## 3\. 单元测试', r'^## 4\. 集成测试'))
    merged.append('')
    merged.append('### 6.3 集成测试\n')
    merged += shift_headings(get_section_text(test, r'^## 4\. 集成测试', r'^## 5\. 确认测试'))
    merged.append('')
    merged.append('### 6.4 确认测试\n')
    merged += shift_headings(get_section_text(test, r'^## 5\. 确认测试', r'^## 6\. 黑盒测试'))
    merged.append('')
    merged.append('### 6.5 黑盒测试\n')
    merged += shift_headings(get_section_text(test, r'^## 6\. 黑盒测试', r'^## 7\. 系统测试'))
    merged.append('')
    merged.append('### 6.6 系统测试\n')
    merged += shift_headings(get_section_text(test, r'^## 7\. 系统测试', r'^## 8\. 测试总结'))
    merged.append('')
    merged.append('### 6.7 测试总结\n')
    merged += shift_headings(get_section_text(test, r'^## 8\. 测试总结', r'^---'))
    
    merged_text = '\n'.join(merged)
    merged_text = convert_bullet_sections(merged_text)
    return merged_text

# ============================================================
# 4. DOCX 生成（与之前相同：三线表 + 南信大格式）
# ============================================================
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def apply_three_line_table(table):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"},
                    bottom={"val": "single", "sz": "6", "color": "000000"},
                    left={"val": "nil"}, right={"val": "nil"})
            elif i == len(table.rows) - 1:
                set_cell_border(cell, top={"val": "nil"},
                    bottom={"val": "single", "sz": "12", "color": "000000"},
                    left={"val": "nil"}, right={"val": "nil"})
            else:
                set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"},
                    left={"val": "nil"}, right={"val": "nil"})

def add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                            size=Pt(12), bold=False, alignment=None, first_line_indent=None,
                            space_before=0, space_after=0, line_spacing=1.5):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    
    # 解析 **...** 行内加粗：拆分为普通文本和加粗文本交替
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = size
        run.font.name = font_name_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
        # 奇数索引（1,3,5...）是捕获组内容，即加粗文本
        if idx % 2 == 1:
            run.font.bold = True
        else:
            run.font.bold = bold
    return p

def add_table_caption(doc, text):
    return add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                                   size=Pt(10.5), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                   space_before=6, space_after=3)

def parse_md_table(lines, start_idx):
    if start_idx >= len(lines):
        return None, None, start_idx
    header_line = lines[start_idx].strip()
    if not header_line.startswith('|') or not header_line.endswith('|'):
        return None, None, start_idx
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    sep_idx = start_idx + 1
    if sep_idx < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[sep_idx].strip()):
        sep_idx += 1
    else:
        return None, None, start_idx
    rows = []
    i = sep_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
            i += 1
        else:
            break
    return headers, rows, i

def build_docx(merged_text, output_name):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    
    lines = merged_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
        
        if stripped == '---':
            i += 1
            continue
        
        # 一级标题 #
        if re.match(r'^#\s+', stripped):
            title = re.sub(r'^#\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=12, space_after=6)
            i += 1
            continue
        
        # 二级标题 ##
        if re.match(r'^##\s+', stripped):
            title = re.sub(r'^##\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(14), bold=True, space_before=12, space_after=3)
            i += 1
            continue
        
        # 三级标题 ###
        if re.match(r'^###\s+', stripped):
            title = re.sub(r'^###\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=6, space_after=3)
            i += 1
            continue
        
        # 四级+标题
        if re.match(r'^#{4,}\s+', stripped):
            title = re.sub(r'^#{4,}\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=3, space_after=3)
            i += 1
            continue
        
        # 引用
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                                    size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        
        # MD 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            # 表题（如 **表1-1**）
            if re.match(r'^\|\s*\*\*.*\*\*\s*\|$', stripped):
                caption = stripped.strip('|').strip().strip('*').strip()
                add_table_caption(doc, caption)
                i += 1
                if i < len(lines):
                    next_line = lines[i].strip()
                    if next_line.startswith('|') and next_line.endswith('|') and not re.match(r'^\|[\s\-:|]+\|$', next_line):
                        headers, rows, i = parse_md_table(lines, i)
                        if headers:
                            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            for j, h in enumerate(headers):
                                cell = table.rows[0].cells[j]
                                cell.text = h
                                for p in cell.paragraphs:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in p.runs:
                                        run.font.size = Pt(10.5)
                                        run.font.bold = True
                                        run.font.name = 'Times New Roman'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            ncols = len(headers)
                            for ri, row in enumerate(rows):
                                if len(row) < ncols:
                                    row += [''] * (ncols - len(row))
                                elif len(row) > ncols:
                                    row = row[:ncols]
                                for j, cell_text in enumerate(row):
                                    cell = table.rows[ri+1].cells[j]
                                    cell.text = cell_text
                                    for p in cell.paragraphs:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in p.runs:
                                            run.font.size = Pt(10.5)
                                            run.font.name = 'Times New Roman'
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            apply_three_line_table(table)
                            doc.add_paragraph()
                            continue
                continue
            
            caption_candidate = None
            check_i = i - 1
            while check_i >= 0 and not lines[check_i].strip():
                check_i -= 1
            if check_i >= 0:
                prev = lines[check_i].strip()
                if prev.startswith('**表') and '**' in prev[3:]:
                    caption_candidate = prev.strip('*').strip()
            
            headers, rows, new_i = parse_md_table(lines, i)
            if headers:
                if caption_candidate:
                    add_table_caption(doc, caption_candidate)
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = h
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                            run.font.bold = True
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                ncols = len(headers)
                for ri, row in enumerate(rows):
                    if len(row) < ncols:
                        row += [''] * (ncols - len(row))
                    elif len(row) > ncols:
                        row = row[:ncols]
                    for j, cell_text in enumerate(row):
                        cell = table.rows[ri+1].cells[j]
                        cell.text = cell_text
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.size = Pt(10.5)
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                apply_three_line_table(table)
                doc.add_paragraph()
                i = new_i
                continue
            else:
                i += 1
                continue
        
        # 代码块
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.size = Pt(9)
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            continue
        
        # 图片
        if stripped.startswith('!['):
            i += 1
            continue
        
        # 独立的 **表...** 行：跳过，由后续表格解析时作为表题使用
        if re.match(r'^\*\*表.*\*\*$', stripped):
            i += 1
            continue
        
        # 普通段落
        # 独立的 **xxx** 行（整行加粗标签）：无缩进
        if re.match(r'^\*\*.*\*\*$', stripped):
            inner = stripped.strip('*').strip()
            add_paragraph_with_font(doc, inner, font_name_cn='宋体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=3, space_after=3)
            i += 1
            continue
        
        add_paragraph_with_font(doc, stripped, font_name_cn='宋体', font_name_en='Times New Roman',
                                size=Pt(12), first_line_indent=Cm(0.74), space_before=0, space_after=0)
        i += 1
    
    output_path = os.path.join(BASE, output_name)
    doc.save(output_path)
    print(f'DOCX 已保存: {output_path}')
    return output_path

# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    print('正在按新结构重组文档...')
    merged = build_restructured_md()
    
    md_output = os.path.join(BASE, '南信猫友记-软件设计说明书.md')
    with open(md_output, 'w', encoding='utf-8') as f:
        f.write(merged)
    print(f'MD 已保存: {md_output}')
    
    print('正在生成 DOCX...')
    build_docx(merged, '南信猫友记-软件设计说明书.docx')
    print('完成！')