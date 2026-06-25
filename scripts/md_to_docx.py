"""将更新后的 MD 文档转换为 DOCX"""
import re
import os
import tempfile
import hashlib
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    import cairosvg
    HAS_CAIRO = True
except:
    HAS_CAIRO = False

BASE = r"C:\Users\16570\Desktop\nanxin-catbook\docs"
DOCX_DIR = os.path.join(BASE, "docx")
CACHE_DIR = os.path.join(BASE, ".img_cache")

FILES = [
    ("需求分析/需求规格说明书.md", "需求规格说明书.docx"),
    ("概要设计/概要设计说明书.md", "概要设计说明书.docx"),
    ("详细设计/详细设计说明书.md", "详细设计说明书.docx"),
    ("测试/测试报告.md", "测试报告.docx"),
]

def resolve_image(img_path, md_path):
    """解析图片路径，返回实际文件路径（如果是SVG则转为PNG）"""
    # 从 MD 文件目录和 BASE 两个角度尝试
    full_path = os.path.normpath(os.path.join(BASE, img_path))
    alt_path = os.path.normpath(os.path.join(os.path.dirname(md_path), img_path))
    actual = full_path if os.path.exists(full_path) else alt_path

    if not os.path.exists(actual):
        return None

    # SVG → PNG 转换
    if actual.lower().endswith('.svg'):
        if not HAS_CAIRO:
            return None  # 无法转换
        os.makedirs(CACHE_DIR, exist_ok=True)
        # 用文件内容 hash 做缓存名
        with open(actual, 'rb') as f:
            raw = f.read()
        h = hashlib.md5(raw).hexdigest()[:12]
        png_path = os.path.join(CACHE_DIR, h + '.png')
        if not os.path.exists(png_path):
            cairosvg.svg2png(bytestring=raw, write_to=png_path, scale=2.0)
        return png_path

    return actual

def md_to_docx(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    
    # ===== 页面设置：A4, 2.5cm边距 =====
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 默认字体：宋体/ Times New Roman 小四 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)  # 小四号 = 12pt
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 1.5倍行距
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24)  # 首行缩进2字符

    # ===== 标题样式 =====
    # 一级标题：三号黑体居中
    # 二级标题：四号黑体左顶格
    # 三级标题：小四号黑体左顶格
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hf = hs.font
        hf.name = 'Times New Roman'
        hf.bold = True
        hf.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            hf.size = Pt(16)  # 三号
            hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            hf.size = Pt(14)  # 四号
        elif level == 3:
            hf.size = Pt(12)  # 小四
        else:
            hf.size = Pt(12)
        hs.paragraph_format.first_line_indent = Pt(0)

    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.startswith("```"):
            if in_code_block:
                # 结束代码块
                in_code_block = False
                # 写为等宽字体段落
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
                code_buffer = []
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line.rstrip('\n'))
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 水平线
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # add a bottom border to simulate hr
            p.paragraph_format.border = None
            # just pass, add spacing
            i += 1
            continue

        # 表格（以 | 开头且包含 --- 分隔行的前一行是表头）
        if line.strip().startswith("|") and "|" in line:
            rows = []
            header_line = line
            i += 1
            # 跳过分隔行
            if i < len(lines) and re.match(r'^\|[\s\-:|]+\|?$', lines[i]):
                i += 1
            # 收集数据行
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            
            # 解析表头
            headers = [h.strip() for h in header_line.strip().strip('|').split('|')]
            # 解析数据
            data_rows = []
            for row in rows:
                cells = [c.strip() for c in row.strip().strip('|').split('|')]
                data_rows.append(cells)
            
            # 创建表格
            table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # 自动调整列宽：根据各列内容最大长度分配比例
            col_max_lens = [len(h) for h in headers]
            for row_data in data_rows:
                for j, cell_text in enumerate(row_data):
                    if j < len(headers):
                        col_max_lens[j] = max(col_max_lens[j], len(cell_text))
            total = sum(col_max_lens) or 1
            table.autofit = False
            avail_cm = 16.0  # A4 21cm - 左右边距2.5cm*2 = 16cm
            for j, clen in enumerate(col_max_lens):
                col_width_cm = avail_cm * clen / total
                for row in table.rows:
                    row.cells[j].width = Cm(col_width_cm)
            
            # 填充表头
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 填充数据
            for ri, row_data in enumerate(data_rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(headers):
                        cell = table.rows[ri + 1].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Times New Roman'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()  # 表后间距
            continue

        # 标题
        heading_match = re.match(r'^#{1,4}\s+(.+)$', line)
        if heading_match:
            level = len(re.match(r'^(#+)', line).group(1))
            text = heading_match.group(1)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 列表项
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if not list_match:
            list_match = re.match(r'^(\s*)\d+[\.)]\s+(.+)$', line)
        if list_match:
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            # 处理粗体
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # 图片
        img_match = re.match(r'^!\[.*?\]\((.+?)\)$', line)
        if img_match:
            img_path = img_match.group(1)
            actual = resolve_image(img_path, md_path)
            if actual and os.path.exists(actual):
                try:
                    doc.add_picture(actual, width=Inches(5.2))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    p = doc.add_paragraph(f"[图片: {img_path}]")
            else:
                p = doc.add_paragraph(f"[图片: {img_path}]")
            i += 1
            continue

        # 普通段落（含粗体、行内代码）
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            # 表题/图题居中
            if text.startswith('**表') or text.startswith('**图'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)

            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                else:
                    run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        i += 1

    doc.save(docx_path)
    print(f"  OK -> {docx_path}")

if __name__ == "__main__":
    for md_rel, docx_name in FILES:
        md_full = os.path.join(BASE, md_rel)
        docx_full = os.path.join(DOCX_DIR, docx_name)
        print(f"Converting: {md_rel} → {docx_name}")
        if os.path.exists(md_full):
            md_to_docx(md_full, docx_full)
        else:
            print(f"  MISSING: {md_full}")
    print("Done!")
