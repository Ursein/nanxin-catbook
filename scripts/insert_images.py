"""
Insert screenshots into the detailed design docx, replacing ASCII diagrams.
"""
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


def insert_images(docx_path, img_dir):
    doc = Document(docx_path)

    # Image config: (heading text to find after, image path, img width)
    inserts = [
        ("4.3 首页界面设计", os.path.join(img_dir, "主页.png"), Inches(5.5)),
        ("4.4 详情页界面设计", os.path.join(img_dir, "猫咪详情.png"), Inches(5.5)),
    ]

    for heading_text, img_path, width in inserts:
        if not os.path.exists(img_path):
            print(f"[WARN] Image not found: {img_path}")
            continue

        # Find the heading paragraph
        heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading') and heading_text in p.text:
                heading_idx = i
                break

        if heading_idx is None:
            print(f"[WARN] Heading not found: {heading_text}")
            continue

        # Find the "布局结构：" paragraph after the heading
        layout_idx = None
        for i in range(heading_idx + 1, min(heading_idx + 10, len(doc.paragraphs))):
            if "布局结构" in doc.paragraphs[i].text:
                layout_idx = i
                break

        if layout_idx is None:
            print(f"[WARN] '布局结构' not found after {heading_text}")
            continue

        # Find the ASCII code block (```) content that follows
        # The ASCII art is in the paragraphs after 布局结构, up to the 说明： paragraph
        # We need to find where the ASCII art ends and insert the image after it
        # The ASCII art paragraphs are between 布局结构： (which has the code fence) and 说明：
        # Let's find the "说明：" paragraph or the next heading
        end_idx = None
        for i in range(layout_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if text.startswith('说明') or text.startswith('**说明') or doc.paragraphs[i].style.name.startswith('Heading'):
                end_idx = i
                break

        if end_idx is None:
            end_idx = layout_idx + 20  # fallback

        # Mark paragraphs to clear (the ASCII diagram content)
        clear_indices = list(range(layout_idx + 1, end_idx))

        # Insert image after the layout_idx paragraph
        # We'll clear the ASCII content paragraphs and put the image there
        if clear_indices:
            # Reuse the first paragraph for the image
            img_paragraph = doc.paragraphs[clear_indices[0]]
            img_paragraph.clear()

            # Add image centered
            run = img_paragraph.add_run()
            run.add_picture(img_path, width=width)
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add a caption paragraph after the image
            caption_text = f"图：{os.path.splitext(os.path.basename(img_path))[0]} 实际界面"
            # Use the next paragraph or add a new one
            if len(clear_indices) > 1:
                cap_p = doc.paragraphs[clear_indices[1]]
                cap_p.clear()
            else:
                cap_p = doc.add_paragraph()

            cap_run = cap_p.add_run(caption_text)
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Clear remaining ASCII paragraphs
            for idx in clear_indices[2:]:
                p = doc.paragraphs[idx]
                p.clear()

        print(f"[OK] Inserted {os.path.basename(img_path)} after {heading_text}")

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: py scripts/insert_images.py <docx_path> <img_dir>")
        sys.exit(1)

    insert_images(sys.argv[1], sys.argv[2])