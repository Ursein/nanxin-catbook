"""
南信猫友记 - 软件设计文档生成器
合并四份 MD → 自然段改写 → 生成 DOCX（三线表 + 南信大毕业论文格式）
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# 1. 读取四份 MD 文件
# ============================================================
BASE = r'C:\Users\16570\Desktop\nanxin-catbook\docs'
MD_FILES = [
    os.path.join(BASE, '需求分析', '需求规格说明书.md'),
    os.path.join(BASE, '概要设计', '概要设计说明书.md'),
    os.path.join(BASE, '详细设计', '详细设计说明书.md'),
    os.path.join(BASE, '测试', '测试报告.md'),
]

def read_md(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# ============================================================
# 2. 自然段改写：短句列表 → 自然段
# ============================================================
def bullets_to_paragraph(lines, category_prefix=""):
    """将短句列表改写为自然段"""
    sentences = []
    for line in lines:
        text = line.strip().lstrip('-').lstrip('•').strip()
        if text:
            # 去掉末尾句号，后续统一加
            text = text.rstrip('。').rstrip('.').strip()
            sentences.append(text)
    if not sentences:
        return ""
    if not category_prefix:
        category_prefix = ""
    if len(sentences) == 1:
        return f"{category_prefix}{sentences[0]}。"
    if len(sentences) == 2:
        return f"{category_prefix}{sentences[0]}，{sentences[1]}。"
    result = f"{category_prefix}{sentences[0]}，{'，'.join(sentences[1:-1])}，{sentences[-1]}。"
    return result

def convert_bullet_sections(text):
    """识别并转换短句列表为自然段"""
    lines = text.split('\n')
    result = []
    i = 0
    in_bullet_group = False
    bullet_lines = []
    bullet_context = ""
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 检测"安全性需求"等标题后紧跟的 bullet list
        if re.match(r'^#{1,4}\s+\S', stripped):
            if in_bullet_group and bullet_lines:
                para = bullets_to_paragraph(bullet_lines, bullet_context)
                result.append('')
                result.append(para)
                result.append('')
                bullet_lines = []
                in_bullet_group = False
            result.append(line)
            bullet_context = ""
            i += 1
            continue
        
        # 检测自然语言标题（如 "#### 安全性需求"）
        if re.match(r'^#{3,5}\s*(安全性|可用性|可维护性|性能|约束|条件|限制|结论)', stripped):
            if in_bullet_group and bullet_lines:
                para = bullets_to_paragraph(bullet_lines, bullet_context)
                result.append('')
                result.append(para)
                result.append('')
                bullet_lines = []
            result.append(line)
            bullet_context = ""
            # 提取标题文本作为自然段前缀
            title_text = re.sub(r'^#+\s*', '', stripped)
            if title_text.endswith('需求') or title_text.endswith('性'):
                bullet_context = f"在{title_text}方面，"
            elif title_text:
                bullet_context = f"{title_text}方面，"
            in_bullet_group = False
            i += 1
            continue
        
        # 检测 bullet 行
        if re.match(r'^[\s]*[-•]\s+', stripped):
            if not in_bullet_group:
                in_bullet_group = True
            bullet_lines.append(stripped)
            i += 1
            continue
        
        # 非 bullet 行
        if in_bullet_group and bullet_lines:
            para = bullets_to_paragraph(bullet_lines, bullet_context)
            result.append('')
            result.append(para)
            result.append('')
            bullet_lines = []
            in_bullet_group = False
            bullet_context = ""
        
        result.append(line)
        i += 1
    
    # 处理文件末尾的 bullet group
    if in_bullet_group and bullet_lines:
        para = bullets_to_paragraph(bullet_lines, bullet_context)
        result.append('')
        result.append(para)
        result.append('')
    
    return '\n'.join(result)

# ============================================================
# 3. 合并四份 MD
# ============================================================
def merge_md_files():
    """合并四份 MD 文件，调整标题层级"""
    merged = []
    merged.append('# 南信猫友记——软件设计文档\n')
    merged.append('> 南京信息工程大学 计算机与软件学院\n')
    merged.append('> 本文档整合了需求规格说明书、概要设计说明书、详细设计说明书及测试报告。\n')
    merged.append('\n---\n')
    
    section_titles = [
        '第一部分：需求规格说明书',
        '第二部分：概要设计说明书',
        '第三部分：详细设计说明书',
        '第四部分：测试报告',
    ]
    
    for idx, (path, title) in enumerate(zip(MD_FILES, section_titles)):
        content = read_md(path)
        # 去掉原有的 # 一级标题行
        content = re.sub(r'^#\s+.*\n', '', content, count=1)
        
        # 标题层级下调一级：## → ###, ### → ####, #### → #####
        content = re.sub(r'^####\s', '##### ', content, flags=re.MULTILINE)
        content = re.sub(r'^###\s', '#### ', content, flags=re.MULTILINE)
        content = re.sub(r'^##\s', '### ', content, flags=re.MULTILINE)
        
        # 添加分部标题
        merged.append(f'# {title}\n')
        merged.append(content)
        if idx < len(MD_FILES) - 1:
            merged.append('\n---\n')
    
    merged_text = '\n'.join(merged)
    # 自然段改写
    merged_text = convert_bullet_sections(merged_text)
    return merged_text

# ============================================================
# 4. 生成 DOCX（南信大毕业论文格式 + 三线表）
# ============================================================
def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def apply_three_line_table(table):
    """对表格应用三线表样式：顶线粗、表头底线细、底线粗，无竖线"""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                # 表头行：顶线粗(1.5pt=12)、底线细(0.75pt=6)
                set_cell_border(cell,
                    top={"val": "single", "sz": "12", "color": "000000"},
                    bottom={"val": "single", "sz": "6", "color": "000000"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            elif i == len(table.rows) - 1:
                # 最后一行：底线粗(1.5pt=12)
                set_cell_border(cell,
                    top={"val": "nil"},
                    bottom={"val": "single", "sz": "12", "color": "000000"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            else:
                # 中间行：无边框
                set_cell_border(cell,
                    top={"val": "nil"},
                    bottom={"val": "nil"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )

def add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                            size=Pt(12), bold=False, alignment=None, first_line_indent=None,
                            space_before=0, space_after=0, line_spacing=1.5):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    
    # 解析 **...** 行内加粗：拆分为普通文本和加粗文本交替
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = size
        run.font.name = font_name_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
        if idx % 2 == 1:
            run.font.bold = True
        else:
            run.font.bold = bold
    return p

def add_table_caption(doc, text):
    """添加表题（表上方居中，宋体五号）"""
    p = add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                                size=Pt(10.5), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=6, space_after=3)
    return p

def parse_md_table(lines, start_idx):
    """解析 MD 表格，返回 (headers, rows, end_idx)"""
    if start_idx >= len(lines):
        return None, None, start_idx
    
    header_line = lines[start_idx].strip()
    if not header_line.startswith('|') or not header_line.endswith('|'):
        return None, None, start_idx
    
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    
    # 跳过分隔行
    sep_idx = start_idx + 1
    if sep_idx < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[sep_idx].strip()):
        sep_idx += 1
    else:
        return None, None, start_idx
    
    rows = []
    i = sep_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
            i += 1
        else:
            break
    
    return headers, rows, i

def build_docx(merged_text):
    """根据合并后的 MD 文本生成 DOCX"""
    doc = Document()
    
    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)  # 小四
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    
    lines = merged_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 跳过空行
        if not stripped:
            i += 1
            continue
        
        # 分隔线
        if stripped == '---':
            i += 1
            continue
        
        # 一级标题 #
        if re.match(r'^#\s+', stripped):
            title = re.sub(r'^#\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=12, space_after=6)
            i += 1
            continue
        
        # 二级标题 ##
        if re.match(r'^##\s+', stripped):
            title = re.sub(r'^##\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(14), bold=True, space_before=12, space_after=3)
            i += 1
            continue
        
        # 三级标题 ###
        if re.match(r'^###\s+', stripped):
            title = re.sub(r'^###\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=6, space_after=3)
            i += 1
            continue
        
        # 四级标题 ####
        if re.match(r'^####\s+', stripped):
            title = re.sub(r'^####\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=6, space_after=3)
            i += 1
            continue
        
        # 五级标题 #####
        if re.match(r'^#####\s+', stripped):
            title = re.sub(r'^#####\s+', '', stripped)
            add_paragraph_with_font(doc, title, font_name_cn='黑体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=3, space_after=3)
            i += 1
            continue
        
        # 引用 >
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            add_paragraph_with_font(doc, text, font_name_cn='宋体', font_name_en='Times New Roman',
                                    size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        
        # MD 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            # 可能是表题（如 **表1-1 国内外同类系统对比**）
            if re.match(r'^\|\s*\*\*.*\*\*\s*\|$', stripped):
                caption = stripped.strip('|').strip().strip('*').strip()
                add_table_caption(doc, caption)
                i += 1
                # 检查下一行是否也是表格（可能是表题单独一行）
                if i < len(lines) and lines[i].strip().startswith('|') and re.match(r'^\|[\s\-:|]+\|$', lines[i].strip().lstrip('|')):
                    # 这是表头行，继续解析
                    pass
                else:
                    # 检查是否下一行是表头
                    if i < len(lines):
                        next_line = lines[i].strip()
                        if next_line.startswith('|') and next_line.endswith('|') and not re.match(r'^\|[\s\-:|]+\|$', next_line):
                            headers, rows, i = parse_md_table(lines, i)
                            if headers:
                                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                # 表头
                                for j, h in enumerate(headers):
                                    cell = table.rows[0].cells[j]
                                    cell.text = h
                                    for p in cell.paragraphs:
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in p.runs:
                                            run.font.size = Pt(10.5)
                                            run.font.bold = True
                                            run.font.name = 'Times New Roman'
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                # 数据行
                                ncols = len(headers)
                                for ri, row in enumerate(rows):
                                    if len(row) < ncols:
                                        row += [''] * (ncols - len(row))
                                    elif len(row) > ncols:
                                        row = row[:ncols]
                                    for j, cell_text in enumerate(row):
                                        cell = table.rows[ri+1].cells[j]
                                        cell.text = cell_text
                                        for p in cell.paragraphs:
                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in p.runs:
                                                run.font.size = Pt(10.5)
                                                run.font.name = 'Times New Roman'
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                apply_three_line_table(table)
                                doc.add_paragraph()  # 表后空行
                                continue
                continue
            
            # 正式解析表格
            caption_candidate = None
            # 看前一行是不是表题
            check_i = i - 1
            while check_i >= 0 and not lines[check_i].strip():
                check_i -= 1
            if check_i >= 0:
                prev = lines[check_i].strip()
                if prev.startswith('**表') and '**' in prev[3:]:
                    caption_candidate = prev.strip('*').strip()
            
            headers, rows, new_i = parse_md_table(lines, i)
            if headers:
                if caption_candidate:
                    add_table_caption(doc, caption_candidate)
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 表头
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = h
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(10.5)
                            run.font.bold = True
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                # 数据行
                ncols = len(headers)
                for ri, row in enumerate(rows):
                    if len(row) < ncols:
                        row += [''] * (ncols - len(row))
                    elif len(row) > ncols:
                        row = row[:ncols]
                    for j, cell_text in enumerate(row):
                        cell = table.rows[ri+1].cells[j]
                        cell.text = cell_text
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.size = Pt(10.5)
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                apply_three_line_table(table)
                doc.add_paragraph()
                i = new_i
                continue
            else:
                i += 1
                continue
        
        # 代码块
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.size = Pt(9)
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            continue
        
        # 图片引用
        if stripped.startswith('!['):
            i += 1
            continue
        
        # 独立的 **表...** 行：跳过，由后续表格解析时作为表题使用
        if re.match(r'^\*\*表.*\*\*$', stripped):
            i += 1
            continue
        
        # 普通段落
        # 独立的 **xxx** 行（整行加粗标签）：无缩进
        if re.match(r'^\*\*.*\*\*$', stripped):
            inner = stripped.strip('*').strip()
            add_paragraph_with_font(doc, inner, font_name_cn='宋体', font_name_en='Times New Roman',
                                    size=Pt(12), bold=True, space_before=3, space_after=3)
            i += 1
            continue
        
        add_paragraph_with_font(doc, stripped, font_name_cn='宋体', font_name_en='Times New Roman',
                                size=Pt(12), first_line_indent=Cm(0.74), space_before=0, space_after=0)
        i += 1
    
    # 保存
    output_path = os.path.join(BASE, '南信猫友记-软件设计文档.docx')
    doc.save(output_path)
    print(f'DOCX 已保存到: {output_path}')
    return output_path

# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    print('正在合并 Markdown 文件...')
    merged = merge_md_files()
    
    md_output = os.path.join(BASE, '南信猫友记-软件设计文档.md')
    with open(md_output, 'w', encoding='utf-8') as f:
        f.write(merged)
    print(f'合并 MD 已保存到: {md_output}')
    
    print('正在生成 DOCX（三线表 + 南信大毕业论文格式）...')
    build_docx(merged)
    print('完成！')